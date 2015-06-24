# -*- coding: utf-8 -*-
"""
Created on Tue Jun 16 15:52:53 2015

@author: chris sewell
"""
from math import log10, floor
from io import BytesIO
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Cm

from pandas import DataFrame
import numpy as np

class MSDocument(object):
    """a class to output a Microsoft Word Document
    
    NB: docx.api.Document can't be directly inherited as it is a function which 
    returns various classes dependent on the *docx* parameter 
    """
    def __init__(self, docx=None):
        """a class to output a Microsoft Word Document

        inherited api details for :py:class:`docx.document.Document` can be 
        found at; https://python-docx.readthedocs.org/en/latest/api/document.html 
        
        the class has an internal state for the number of calls to add_picture
        and add_table for use in caption numbering
        
        Parameters
        ----------
        docx : str or file-like object
            can be either a path to a .docx file (a string) or a file-like object. 
            If docx is missing or None, the built-in default document “template” 
            is loaded.
        
        """
        self._docx = Document(docx=docx)
        
        self._piccount = 0
        self._tablecount = 0
    
    def __getattr__(self, name):
        """ required to get :py:class:`docx.document.Document` methods """
        return getattr(self._docx, name)
    
    def __dir__(self):
        """ required to have :py:class:`docx.document.Document` methods in 
        :py:mod:`IPython` tab completion"""
        dirlist = self.__class__.__dict__.keys() + self._docx.__class__.__dict__.keys()
        return sorted(dirlist)           

    def add_picture(self, image_path_or_stream, width=None, height=None):
        """ Return a new picture shape added in its own paragraph at the end of
        the document. The picture contains the image at
        *image_path_or_stream*, scaled based on *width* and *height*. If
        neither width nor height is specified, the picture appears at its
        native size. If only one is specified, it is used to compute
        a scaling factor that is then applied to the unspecified dimension,
        preserving the aspect ratio of the image. The native size of the
        picture is calculated using the dots-per-inch (dpi) value specified
        in the image file, defaulting to 72 dpi if no value is specified, as
        is often the case.
        """
        self._piccount += 1
        return self._docx.add_picture(image_path_or_stream, width, height)
    
    def add_table(self, rows, cols, style=None):
        """Add a table having row and column counts of *rows* and *cols*
        respectively and table style of *style*. *style* may be a paragraph
        style object or a paragraph style name. If *style* is |None|, the
        table inherits the default table style of the document.
        """
        self._tablecount += 1
        return self._docx.add_table(rows, cols, style)

    _MARKUPS = {
            'italic':('*','*'),
            'bold':('**', '**'),
            'subscript':('_{', '}'),
            'superscript':('^{', '}'),
            'strike':('~~','~~'),
            'math': ('$', '$')
            }

    def _get_markup(self, para, markup_dict=None):
        """get markup """
        if not markup_dict:
            markup_dict = self._MARKUPS
            
        df = DataFrame(markup_dict, index=['Enter', 'Exit']).T
        df['In']=False
        
        sects=[]
        place=0
        while place > -1:
            place = -1
            markup = None
            estr = None
            for mark, enter in df[df.In==False].Enter.iterkv():
                find = para.find(enter)
                if find > -1 and (find<=place or place==-1):
                    if find == place and len(enter) < len(estr):
                        continue
                    place = find
                    markup = mark
                    estr = enter
            for mark, exit in df[df.In==True].Exit.iterkv():
                find = para.find(exit)
                if find > -1 and (find<=place or place==-1):
                    if find == place and len(exit) < len(estr):
                        continue
                    place = find
                    markup = mark
                    estr = exit
        
            if place > -1:
                sects.append([para[:place], df[df.In==True].index.tolist()])
                df.set_value(markup, 'In', not df.get_value(markup, 'In'))
                para = para[place+len(estr):]

        if df.In.any():
            raise ValueError(
                'the markup does not exit from;\n{}'.format(df[df.In==True]))
            
        sects.append([para, []])
                         
        return sects

    def add_markdown(self, text='', style='Body Text', 
                     markup_dict=None, para=None):
        r"""adds a paragraph to the document, allowing for
        paragraph/font styling akin to a stripped down version of
        markdown text:
        
        paragraph level::
        
            # Header (level denoted by number of #'s)            
            - bullet list
            1. numbered list  
        
        font level::
        
            **bold** 
            *italic* 
            _{subscript} 
            ^{superscript} 
            ~~strikethrough~~ 
            $mathML$
        
        Parameters
        ----------
        text : str
            the text to add
        style : str
            the style to apply (overriden if paragraph level markdown)
        markup_dict : dict
            if set will override built in font level markup
            {font_attribute:(start_chars, end_chars)}
        para : docx.text.paragraph.Paragraph
            a pre-existing paragraph to add the text to
            if set, will ignore paragraph level markdown
        
        Returns
        -------
        para : docx.text.paragraph.Paragraph
            a paragraph added to the document
            
        """
        list_pattern = re.compile('^[-+]\s')
        number_pattern = re.compile('^\d+[.]\s')
        head_pattern = re.compile('^#+\s')
        level=0
        
        if re.match(list_pattern, text):
            style = 'List Bullet'
            text = text[len(re.findall(list_pattern, text)[0]):]
        elif re.match(number_pattern, text):
            style = 'List Number'
            text = text[len(re.findall(number_pattern, text)[0]):]
        elif re.match(head_pattern, text):
            level = len(re.findall(head_pattern, text)[0]) - 1
            text = text[level+1:]
 
        if not para:
            if level:
                para = self.add_heading(level=level)
            else:
                para = self.add_paragraph(style=style)
        if not text:
            return para
        
        sects = self._get_markup(text, markup_dict)
        for txt, markups in sects:
            run = para.add_run(txt)
            font = run.font
            for markup in markups:
                setattr(font, markup, True)

        return para


    def _split_special_paras(self, text):
        """split text into paras if a header or list,
        denominated by; # heading, - bullet or 1. numbered
        """
        patterns = ['[-+]', '\d+[.]', '#+']

        for pattern in patterns:
            if re.match(re.compile('^{}\s'.format(pattern)), text):
                starts = re.findall(re.compile('\n\s*{}\s'.format(pattern)), '\n'+text)
                texts = re.split(re.compile('\n\s*{}\s'.format(pattern)), '\n'+text)
                return [s[1:]+t for s, t in zip(starts, texts[1:])]
        
        return [text]
        

    def add_docstring(self, docstring, style='Body Text',
                      markdown=True):
        """adds a doctring to the document
            
        this function will split text into paragraphs 
        (denominated by a separating blank line)
        remove new-line characters and add to document, allowing for 
        markdown style text designated in 
        :py:func:`pygauss.docs.MSDocument.add_markdown`
        
        Parameters
        ----------
        text : str
            the text to add
        style : str
            the style to apply for each paragraph
        markdown : bool
            whether to apply markdown to the text
        
        Returns
        -------
        paras : docx.text.paragraph.Paragraph
            a list of paragraphs added to the document        

        """
        docx_paras = []
        para_pattern = re.compile('\n[\s]*\n')
        paras = re.split(para_pattern, docstring)

        # remove initial linespace if present
        if paras[0][:1] == '\n':
            paras[0] = paras[0][1:]

        for para in paras:
            if markdown:
                para = para.strip()
                for p in self._split_special_paras(para):
                    p = p.replace('\n', ' ').strip()
                    docx_paras.append(self.add_markdown(p, style=style))
            else:
                para = para.replace('\n', ' ').strip()
                docx_paras.append(self.add_paragraph(para, style=style))
    
        return docx_paras
        
    def add_list(self, text_list=[], numbered=False):
        """adds a list """
        if numbered:
            style='List Number'
        else:
            style='List Bullet'
            
        return [self.add_paragraph(tx, style=style) for tx in text_list]
    
    def add_mpl(self, fig, dpi=None, width=None, height=None, pad_inches=0.2,
                caption=None):
        """add matplotlib figure to the document 
        
        Parameters
        ----------
        fig : matplotlib.figure.Figure
            a matplotlib figure
        dpi : int
            Dots per inch
        width : float
            width of image in document
        height : float
            width of image in document
        pad_inches : float
            amount of padding around the figure
        caption : str
            a caption below the figure

        Returns
        -------
        pic : docx.shape.InlineShape
            an inline picture added to the document        

        """
        stream = BytesIO()
        fig.savefig(stream, format='png', dpi=dpi,
                    bbox_inches='tight', pad_inches=pad_inches,
                    transparent=True)
        
        width = Cm(width) if width else None
        height = Cm(height) if height else None
        
        pic = self.add_picture(stream, width=width, height=height)
        self.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if caption is not None:
            self.add_markdown('***Figure '+ str(self._piccount) + ':*** ' + str(caption),
                              style='Caption')
        
        return pic
           
    def _add_headrw(self, cell, val):
        """ add value to header table cell """
        val = '' if val is None else str(val)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        self.add_markdown('**'+val+'**', para=p)
        return p        
    
    def _sigfigs(self, val, sig_figures=5):
        """round to significant figure"""
        if type(val) == bool:
            return val            
        try:
            if val >= 0.:
                return round(val, -int(floor(log10(val))) + (sig_figures-1))
            else:
                return -round(-val, -int(floor(log10(-val))) + (sig_figures-1))
        except Exception:
            return val

    def add_dataframe(self, df, incl_indx=True, autofit=True, sig_figures=5,
                      style='Medium List 1 Accent 1', caption=None):
        """add dataframe as a table to the document
        
        Parameters
        ----------
        df : pandas.DataFrame
            a pandas dataframe
        incl_indx : bool
            include dataframes index in table
        autofit : bool
            allow table to autofit content 
        sig_figures : int
            number of significant figures for numbers in table
        style : str
            MS Word table style
        caption : str
            add a caption below the table

        Returns
        -------
        pic : docx.table.Table
            a table added to the document        
        
        """
        df = df.fillna('-')
        rows, cols = df.shape 

        if hasattr(df.columns, 'levels'):
            hrows = len(df.columns.levels)
        else:
            hrows = 1
        
        if incl_indx:
            if hasattr(df.index, 'levels'):
                icols = len(df.index.levels)
            else:
                icols = 1
        else:
            icols = 0
            
        table = self.add_table(rows=rows+hrows, cols=cols+icols, 
                                     style=style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = autofit
         
        #add header rows
        if hasattr(df.columns, 'levels'):
            
            h_array = np.array(df.columns.tolist())
            col_count, rw_count = h_array.shape
            
            for rw in range(rw_count):
                #merge adjacent headers with the same value for higher level indexes
                rw_values = h_array[:,rw].tolist()
                start_col = end_col = icols
                while rw_values:
                    val = rw_values.pop(0)
                    if rw_values and rw < (rw_count-1):
                        while val == rw_values[0]:
                            rw_values.pop(0)
                            end_col += 1
                            if not rw_values:
                                break
                    cell = table.rows[rw].cells[start_col]
                    if not start_col == end_col:
                        cell = cell.merge(table.rows[rw].cells[end_col])
                    self._add_headrw(cell, val)
                    start_col = end_col = end_col + 1
                
        else:
            for col, val in enumerate(df.keys()):
                cell = table.rows[hrows-1].cells[col+icols]
                self._add_headrw(cell, val)

        if incl_indx:
            if hasattr(df.index, 'levels'):
               for col, name in enumerate(df.index.names):
                   cell = table.rows[hrows-1].cells[col]
                   self._add_headrw(cell, name)
                   row = hrows-2
                   #ensure all cells are formatted correctly
                   while row >= 0:
                       cell = table.rows[row].cells[col]
                       self._add_headrw(cell, '')
                       row -= 1
            else:
                cell = table.rows[hrows-1].cells[0]
                self._add_headrw(cell, df.index.name)
                row = hrows-2
                #ensure all cells are formatted correctly
                while row >= 0:
                    cell = table.rows[row].cells[0]
                    self._add_headrw(cell, '')
                    row -= 1
                    
        
                        
        #add data rows
        for row, id_series in enumerate(df.iterrows()):

            if incl_indx:
                if hasattr(df.index, 'levels'):
                    for col, val in enumerate(df.index.tolist()[row]):
                        cell = table.rows[row+hrows].cells[col]
                        self._add_headrw(cell, val)
                else:
                    cell = table.rows[row+hrows].cells[0]
                    self._add_headrw(cell, df.index[row])
                    
                    
            for col, item in enumerate(id_series[1].iteritems()):
                cell = table.rows[row+hrows].cells[col+icols]
                cell.text = str(self._sigfigs(item[1], sig_figures))
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True

        if caption is not None:
            self.add_markdown('***Table '+ str(self._tablecount) + ':*** ' + str(caption),
                              style='Caption')

        return table
    

        
        
